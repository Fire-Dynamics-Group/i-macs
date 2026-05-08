"""Tests for report_docx.py — per-batch DOCX report with charts and tables."""

import io
import zipfile
from pathlib import Path

import pytest
from docx import Document

from macs_automation.db import ResultsDB
from macs_automation.report_docx import generate_batch_docx
from macs_automation.tests.conftest import _insert_populated_run


# ─── Fixture: populated DB with batch_id ─────────────────────────────────────

@pytest.fixture
def batch_db(tmp_path):
    """Database with two batches: batch_A (6 runs), batch_B (3 runs), plus an error."""
    db_path = tmp_path / "batch_test.db"
    db = ResultsDB(db_path, check_same_thread=False)

    batch_a = "batch_A_001"
    batch_b = "batch_B_002"
    db.insert_batch(batch_a, mode="sweep", total_expected=6)
    db.insert_batch(batch_b, mode="lhs", total_expected=3)

    # Batch A: 5 passing + 1 failing
    for i in range(5):
        uf = 0.4 + i * 0.1
        qf = 400.0 + i * 50.0
        wp = 30.0 + i * 10.0
        run_id = _insert_populated_run(db, i, uf_max=round(uf, 2), qf=qf, window_percent=wp)
        db.conn.execute("UPDATE runs SET batch_id = ? WHERE id = ?", (batch_a, run_id))
    # 1 failing in batch A
    run_id = _insert_populated_run(db, 5, uf_max=1.2, qf=750.0, window_percent=90.0)
    db.conn.execute("UPDATE runs SET batch_id = ? WHERE id = ?", (batch_a, run_id))

    # Batch B: 2 passing + 1 error
    for i in range(2):
        uf = 0.5 + i * 0.15
        qf = 500.0 + i * 40.0
        wp = 50.0 + i * 15.0
        run_id = _insert_populated_run(db, 6 + i, uf_max=round(uf, 2), qf=qf, window_percent=wp)
        db.conn.execute("UPDATE runs SET batch_id = ? WHERE id = ?", (batch_b, run_id))
    # Error run in batch B
    run_id = _insert_populated_run(db, 8, error="COMError: timeout")
    db.conn.execute("UPDATE runs SET batch_id = ? WHERE id = ?", (batch_b, run_id))

    db.conn.commit()
    yield db
    db.close()


# ─── Tests: Batch-scoped DB queries ─────────────────────────────────────────

class TestBatchDBMethods:
    def test_get_batch_successful_runs_count(self, batch_db):
        runs = batch_db.get_batch_successful_runs("batch_A_001")
        # 5 pass + 1 fail = 6 successful (no error)
        assert len(runs) == 6

    def test_get_batch_successful_runs_only_that_batch(self, batch_db):
        runs_a = batch_db.get_batch_successful_runs("batch_A_001")
        runs_b = batch_db.get_batch_successful_runs("batch_B_002")
        # Batch A: 6 successful, Batch B: 2 successful (1 is error)
        assert len(runs_a) == 6
        assert len(runs_b) == 2

    def test_get_batch_successful_runs_nonexistent(self, batch_db):
        runs = batch_db.get_batch_successful_runs("nonexistent")
        assert runs == []

    def test_get_batch_time_series_column(self, batch_db):
        rows = batch_db.get_batch_time_series_column("batch_A_001", "lofl_temp")
        assert len(rows) > 0
        # All run_ids should belong to batch A
        run_ids = {r[0] for r in rows}
        batch_a_runs = batch_db.get_batch_successful_runs("batch_A_001")
        batch_a_ids = {r["id"] for r in batch_a_runs}
        assert run_ids == batch_a_ids

    def test_get_batch_time_series_column_invalid(self, batch_db):
        with pytest.raises(ValueError):
            batch_db.get_batch_time_series_column("batch_A_001", "DROP TABLE")

    def test_get_batch_stats(self, batch_db):
        stats = batch_db.get_batch_stats("batch_A_001")
        assert stats["total"] == 6
        assert stats["successful"] == 6
        assert stats["errors"] == 0
        assert stats["pass_count"] == 5
        assert stats["fail_count"] == 1

    def test_get_batch_stats_with_errors(self, batch_db):
        stats = batch_db.get_batch_stats("batch_B_002")
        assert stats["total"] == 3
        assert stats["successful"] == 2
        assert stats["errors"] == 1
        assert stats["pass_count"] == 2
        assert stats["fail_count"] == 0


# ─── Tests: DOCX generation ─────────────────────────────────────────────────

class TestGenerateBatchDocx:
    def test_returns_path(self, batch_db):
        path = generate_batch_docx(batch_db, "batch_A_001")
        assert isinstance(path, Path)
        assert path.exists()
        assert path.suffix == ".docx"

    def test_valid_docx(self, batch_db):
        path = generate_batch_docx(batch_db, "batch_A_001")
        doc = Document(path)
        # Should be loadable without error
        assert len(doc.paragraphs) > 0

    def test_contains_title(self, batch_db):
        path = generate_batch_docx(batch_db, "batch_A_001")
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "MACS+" in text
        assert "batch_A_001" in text

    def test_contains_summary_table(self, batch_db):
        path = generate_batch_docx(batch_db, "batch_A_001")
        doc = Document(path)
        # Find a table with summary header columns
        found = False
        for table in doc.tables:
            headers = [cell.text for cell in table.rows[0].cells]
            if "Sim" in headers[0] or "sim" in headers[0].lower():
                found = True
                # Should have header + 6 data rows
                assert len(table.rows) == 7
                break
        assert found, "Summary table not found in DOCX"

    def test_contains_beam_temps_table(self, batch_db):
        path = generate_batch_docx(batch_db, "batch_A_001")
        doc = Document(path)
        found = False
        for table in doc.tables:
            headers = [cell.text for cell in table.rows[0].cells]
            if any("Side A" in h or "Temp A" in h for h in headers):
                found = True
                break
        assert found, "Protected beam temps table not found in DOCX"

    def test_contains_charts(self, batch_db):
        path = generate_batch_docx(batch_db, "batch_A_001")
        doc = Document(path)
        # Charts are embedded as images in the docx
        # Count inline shapes (images)
        image_count = sum(
            1 for rel in doc.part.rels.values()
            if "image" in rel.reltype
        )
        assert image_count == 4, f"Expected 4 chart images, found {image_count}"

    def test_empty_batch_still_generates(self, batch_db):
        """A batch with no runs should still produce a valid DOCX."""
        batch_db.insert_batch("empty_batch", mode="sweep", total_expected=0)
        path = generate_batch_docx(batch_db, "empty_batch")
        doc = Document(path)
        assert len(doc.paragraphs) > 0

    def test_batch_b_has_fewer_rows(self, batch_db):
        path = generate_batch_docx(batch_db, "batch_B_002")
        doc = Document(path)
        # Find summary table — batch B has 2 successful runs
        for table in doc.tables:
            headers = [cell.text for cell in table.rows[0].cells]
            if "sim" in headers[0].lower():
                assert len(table.rows) == 3  # header + 2 data rows
                break


# ─── Tests: API endpoint ─────────────────────────────────────────────────────

@pytest.fixture
def test_client(batch_db, monkeypatch):
    """FastAPI test client pointed at the batch_db."""
    from fastapi.testclient import TestClient
    from macs_automation.app import app
    import macs_automation.app as app_module

    monkeypatch.setattr(app_module, "DB_PATH", Path(batch_db.db_path))
    client = TestClient(app)
    yield client


class TestDocxEndpoint:
    def test_download_with_batch_id(self, test_client):
        resp = test_client.get("/api/report/docx?batch_id=batch_A_001")
        assert resp.status_code == 200
        assert "openxmlformats" in resp.headers["content-type"]
        # Verify it's a valid DOCX (which is a ZIP)
        buf = io.BytesIO(resp.content)
        assert zipfile.is_zipfile(buf)

    def test_download_without_batch_id(self, test_client):
        """Without batch_id, should generate for all runs."""
        resp = test_client.get("/api/report/docx")
        assert resp.status_code == 200
        assert "openxmlformats" in resp.headers["content-type"]

    def test_filename_header(self, test_client):
        resp = test_client.get("/api/report/docx?batch_id=batch_A_001")
        cd = resp.headers.get("content-disposition", "")
        assert "batch_A_001" in cd
        assert ".docx" in cd
