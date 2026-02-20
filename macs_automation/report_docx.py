"""DOCX report generator — per-batch Word document with charts and tables."""

import io
import tempfile
from collections import defaultdict
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

from macs_automation.db import ResultsDB

# ─── Brand styling (matches constants.py / table_constants.py) ───────────────

_FONT_DIR = Path(__file__).parent / "SEGOEUIL.TTF"
_FONT_REGISTERED = False

_LIGHT_TEXT = (0.59, 0.56, 0.56)

_CHART_RC = {
    "xtick.color": _LIGHT_TEXT,
    "ytick.color": _LIGHT_TEXT,
    "axes.titlecolor": _LIGHT_TEXT,
    "axes.labelcolor": _LIGHT_TEXT,
    "axes.edgecolor": _LIGHT_TEXT,
    "legend.labelcolor": _LIGHT_TEXT,
    "figure.figsize": [6, 4],
    "axes.grid": True,
    "grid.linewidth": 0.05,
    "grid.color": _LIGHT_TEXT,
}

_MID_BLUE_20A = "#4798EA33"   # individual runs
_MID_BLUE = "#4798EA"         # scatter pass dots
_CORAL = "coral"              # average line / scatter fail dots


def _register_font():
    """Register Segoe UI Light TTF with matplotlib (once)."""
    global _FONT_REGISTERED
    if _FONT_REGISTERED:
        return
    import matplotlib
    matplotlib.use("Agg")
    from matplotlib import font_manager
    if _FONT_DIR.exists():
        font_manager.fontManager.addfont(str(_FONT_DIR))
        _CHART_RC["font.family"] = "Segoe UI Light"
    _FONT_REGISTERED = True


def _render_timeseries_chart(
    db: ResultsDB,
    column: str,
    ylabel: str,
    runs: list[dict],
    batch_id: Optional[str] = None,
    hline_value: Optional[float] = None,
    legend_loc: str = "center right",
) -> Optional[bytes]:
    """Render a time-series chart to PNG bytes using brand styling."""
    _register_font()
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    if batch_id:
        rows = db.get_batch_time_series_column(batch_id, column)
    else:
        rows = db.get_all_time_series_column(column)

    if not rows:
        return None

    with plt.rc_context(_CHART_RC):
        fig, ax = plt.subplots()

        by_run = defaultdict(list)
        for run_id, time_min, value in rows:
            by_run[run_id].append((time_min, value))

        # Plot individual runs (unlabelled) then one more for the legend
        for rid, data in by_run.items():
            data.sort()
            times = [d[0] for d in data]
            values = [d[1] for d in data]
            ax.plot(times, values, color=_MID_BLUE_20A)
        # Add legend entry via last run
        if by_run:
            last_data = list(by_run.values())[-1]
            ax.plot(
                [d[0] for d in last_data], [d[1] for d in last_data],
                color=_MID_BLUE_20A, label="Recorded Temperature",
            )

        # Compute and plot average
        all_times = set()
        for data in by_run.values():
            for t, _ in data:
                all_times.add(t)
        sorted_times = sorted(all_times)

        if sorted_times and by_run:
            avg_values = []
            for t in sorted_times:
                vals = [dict(data).get(t) for data in by_run.values()]
                vals = [v for v in vals if v is not None]
                avg_values.append(sum(vals) / len(vals) if vals else 0)
            ax.plot(sorted_times, avg_values, color=_CORAL, linewidth=2,
                    label="Average Value")

        if hline_value is not None:
            ax.axhline(y=hline_value, color="red", linewidth=1.5,
                       label="Factored Load")

        ax.set_xlabel("Time (minutes)")
        ax.set_ylabel(ylabel)
        ax.set_ylim(bottom=0)
        # Cut off x-axis at actual sim time (zero to last time, no empty extension)
        if sorted_times:
            ax.set_xlim(left=0, right=sorted_times[-1])
        ax.legend(loc=legend_loc)

        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=300, bbox_inches="tight")
        plt.close(fig)

    buf.seek(0)
    return buf.getvalue()


def _render_scatter_chart(runs: list[dict]) -> Optional[bytes]:
    """Render pass/fail scatter chart to PNG bytes using brand styling."""
    _register_font()
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    pass_qf, pass_wp = [], []
    fail_qf, fail_wp = [], []

    for run in runs:
        qf = run.get("qf")
        wp = run.get("window_percent")
        uf = run.get("uf_max")
        if qf is not None and wp is not None:
            if uf is not None and uf <= 1.0:
                pass_qf.append(qf)
                pass_wp.append(wp)
            else:
                fail_qf.append(qf)
                fail_wp.append(wp)

    if not pass_qf and not fail_qf:
        return None

    with plt.rc_context(_CHART_RC):
        fig, ax = plt.subplots()

        if pass_qf:
            ax.scatter(pass_qf, pass_wp, c=_MID_BLUE, s=5,
                       label="Unity factor < 1.0")
        if fail_qf:
            ax.scatter(fail_qf, fail_wp, c=_CORAL, s=5,
                       label="Unity factor >=1.0")

        ax.set_xlabel("Fireload (MJ/m2)")
        ax.set_ylabel("Glazing Breakage (%)")
        ax.legend()

        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=300, bbox_inches="tight")
        plt.close(fig)

    buf.seek(0)
    return buf.getvalue()


def _fmt(value, decimals=2) -> str:
    """Format a value for table display."""
    if value is None:
        return "\u2014"
    if isinstance(value, float):
        return f"{value:.{decimals}f}"
    return str(value)


def generate_batch_docx(
    db: ResultsDB,
    batch_id: Optional[str] = None,
) -> Path:
    """Generate a DOCX report for a batch (or all runs if batch_id is None).

    Returns the path to the generated .docx file.
    """
    doc = Document()

    # ── Title ────────────────────────────────────────────────────────────
    doc.add_heading("MACS+ Batch Report", level=0)

    if batch_id:
        runs = db.get_batch_successful_runs(batch_id)
        stats = db.get_batch_stats(batch_id)
        doc.add_paragraph(f"Batch ID: {batch_id}")
    else:
        runs = db.get_successful_runs()
        stats = db.get_stats()
        doc.add_paragraph("All runs (no batch filter)")

    # ── Summary statistics ───────────────────────────────────────────────
    doc.add_heading("Summary Statistics", level=1)
    doc.add_paragraph(
        f"Total runs: {stats['total']}  |  "
        f"Successful: {stats['successful']}  |  "
        f"Errors: {stats['errors']}  |  "
        f"Pass: {stats['pass_count']}  |  "
        f"Fail: {stats['fail_count']}"
    )

    # ── Summary table ────────────────────────────────────────────────────
    if runs:
        doc.add_heading("Results Summary", level=1)
        headers = [
            "Sim", "Fire Load\n(MJ/m\u00b2)", "Glazing\n(%)",
            "Max UF", "Time of Max\n(min)", "Time UF>1\n(min)", "Factored Hot\n(kN/m)",
        ]
        table = doc.add_table(rows=1 + len(runs), cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for row_idx, run in enumerate(runs, 1):
            time_of_max = db.get_time_of_max_uf(run["id"])
            time_exceed = db.get_time_exceed_one(run["id"])
            cells = table.rows[row_idx].cells
            cells[0].text = str(row_idx)
            cells[1].text = _fmt(run.get("qf"), 1)
            cells[2].text = _fmt(run.get("window_percent"), 1)
            cells[3].text = _fmt(run.get("uf_max"), 3)
            cells[4].text = _fmt(time_of_max, 1)
            cells[5].text = _fmt(time_exceed, 1)
            cells[6].text = _fmt(run.get("factored_hot"), 2)

        # ── Protected beam critical temps table ──────────────────────────
        doc.add_heading("Protected Beam Critical Temperatures", level=1)
        beam_headers = [
            "Run", "Side A\n(\u00b0C)", "Side B\n(\u00b0C)",
            "Side C\n(\u00b0C)", "Side D\n(\u00b0C)",
            "Fire Load\n(MJ/m\u00b2)", "Glazing\n(%)",
        ]
        beam_table = doc.add_table(rows=1 + len(runs), cols=len(beam_headers))
        beam_table.style = "Light Grid Accent 1"
        beam_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, h in enumerate(beam_headers):
            beam_table.rows[0].cells[i].text = h

        for row_idx, run in enumerate(runs, 1):
            cells = beam_table.rows[row_idx].cells
            cells[0].text = str(row_idx)
            cells[1].text = _fmt(run.get("side_a_critical_temp"), 1)
            cells[2].text = _fmt(run.get("side_b_critical_temp"), 1)
            cells[3].text = _fmt(run.get("side_c_critical_temp"), 1)
            cells[4].text = _fmt(run.get("side_d_critical_temp"), 1)
            cells[5].text = _fmt(run.get("qf"), 1)
            cells[6].text = _fmt(run.get("window_percent"), 1)

    # ── Charts ───────────────────────────────────────────────────────────
    doc.add_heading("Charts", level=1)

    factored_hot = runs[0].get("factored_hot") if runs else None

    # Scatter chart
    scatter_png = _render_scatter_chart(runs)
    if scatter_png:
        doc.add_paragraph("All combinations of Fire Load Density vs Glazing Breakage")
        doc.add_picture(io.BytesIO(scatter_png), width=Inches(5.5))

    # Total capacity vs time
    capacity_png = _render_timeseries_chart(
        db, "total_plate_capacity",
        "Total Slab Capacity (kN/m2)",
        runs, batch_id=batch_id, hline_value=factored_hot,
        legend_loc="center right",
    )
    if capacity_png:
        doc.add_paragraph("Total Capacity Distribution")
        doc.add_picture(io.BytesIO(capacity_png), width=Inches(5.5))

    # Beam temperature vs time
    beam_png = _render_timeseries_chart(
        db, "lofl_temp",
        "Temperature (C)",
        runs, batch_id=batch_id,
        legend_loc="upper right",
    )
    if beam_png:
        doc.add_paragraph("Unprotected Beam Temperature Distribution")
        doc.add_picture(io.BytesIO(beam_png), width=Inches(5.5))

    # Mesh temperature vs time
    mesh_png = _render_timeseries_chart(
        db, "mesh_temp",
        "Temperature (C)",
        runs, batch_id=batch_id,
        legend_loc="upper right",
    )
    if mesh_png:
        doc.add_paragraph("Reinforcement Bar Temperature Distribution")
        doc.add_picture(io.BytesIO(mesh_png), width=Inches(5.5))

    # ── Save ─────────────────────────────────────────────────────────────
    tmp_dir = Path(tempfile.mkdtemp(prefix="macs_docx_"))
    filename = f"macs_report_{batch_id}.docx" if batch_id else "macs_report.docx"
    doc_path = tmp_dir / filename
    doc.save(str(doc_path))
    return doc_path
